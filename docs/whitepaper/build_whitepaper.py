from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT=Path('/mnt/data/PowerChain_Crisis_Capital_Network_Whitepaper_v1.0.0.docx')
LOGO=Path('/mnt/data/logo-green.png')
PWRC=Path('/mnt/data/pwrc.png')
AS=Path('/mnt/data/pc_whitepaper_assets')
GREEN='143C2E'; TEXT='111513'; MUTED='66706A'; BORDER='E3E6E2'; SURFACE='F5F6F4'; BLUE='2457C5'; RED='C9362B'; WHITE='FFFFFF'


def rgb(hexv): return RGBColor.from_string(hexv)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders');tcPr.append(borders)
    for edge in ('top','left','bottom','right'):
        e=borders.find(qn('w:'+edge))
        if e is None: e=OxmlElement('w:'+edge);borders.append(e)
        e.set(qn('w:val'),'single');e.set(qn('w:sz'),size);e.set(qn('w:color'),color)

def margins(cell, top=110, start=135, bottom=110, end=135):
    tcPr=cell._tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar');tcPr.append(mar)
    for name,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        n=mar.find(qn('w:'+name))
        if n is None: n=OxmlElement('w:'+name);mar.append(n)
        n.set(qn('w:w'),str(val));n.set(qn('w:type'),'dxa')

def set_font(run,size=10,color=TEXT,bold=False,name='Liberation Sans'):
    run.font.name=name;run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name);run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size);run.font.color.rgb=rgb(color);run.bold=bold

def page_num(par):
    par.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=par.add_run('POWERCHAIN CRISIS  |  v1.0.0     ');set_font(r,8,MUTED)
    fld=OxmlElement('w:fldSimple');fld.set(qn('w:instr'),'PAGE');par._p.append(fld)

def callout(title,body,tone='green'):
    colors={'green':('F1F6F3',GREEN),'blue':('F1F5FC',BLUE),'red':('FFF3F2',RED)};fill,accent=colors[tone]
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;c=t.cell(0,0);shade(c,fill);border(c,accent,'8');margins(c,140,170,140,170)
    p=c.paragraphs[0];r=p.add_run(title.upper());set_font(r,8,accent,True)
    p2=c.add_paragraph();p2.paragraph_format.space_after=Pt(0);r2=p2.add_run(body);set_font(r2,10,TEXT)
    return t

def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];shade(c,GREEN);border(c,GREEN);margins(c);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r=c.paragraphs[0].add_run(h);set_font(r,8,WHITE,True)
        if widths:c.width=Inches(widths[i])
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            c=cells[i];border(c);margins(c);shade(c,'FFFFFF' if ridx%2==0 else 'FAFBFA');c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            p=c.paragraphs[0];p.paragraph_format.space_after=Pt(0);r=p.add_run(str(v));set_font(r,8.5,TEXT)
            if widths:c.width=Inches(widths[i])
    return t

def eyebrow(text):
    p=doc.add_paragraph(style='Eyebrow');p.add_run(text.upper());return p

def para(text,lead=False):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(7);r=p.add_run(text);set_font(r,12 if lead else 10, MUTED if lead else TEXT);return p

def bullets(items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);r=p.add_run(item);set_font(r,9.5,TEXT)

def figure(filename,caption,width=6.75):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(3);p.add_run().add_picture(str(AS/filename),width=Inches(width))
    c=doc.add_paragraph(caption);c.alignment=WD_ALIGN_PARAGRAPH.CENTER;c.paragraph_format.space_after=Pt(8);r=c.runs[0];set_font(r,8,MUTED);r.italic=True

def newpage(): doc.add_page_break()

def h1(text): return doc.add_heading(text,1)
def h2(text): return doc.add_heading(text,2)

# setup
doc=Document();sec=doc.sections[0]
sec.page_width=Inches(8.27);sec.page_height=Inches(11.69);sec.top_margin=Inches(.68);sec.bottom_margin=Inches(.62);sec.left_margin=Inches(.72);sec.right_margin=Inches(.72);sec.header_distance=Inches(.28);sec.footer_distance=Inches(.28);sec.different_first_page_header_footer=True
styles=doc.styles
normal=styles['Normal'];normal.font.name='Liberation Sans';normal.font.size=Pt(10);normal.font.color.rgb=rgb(TEXT);normal.paragraph_format.line_spacing=1.12;normal.paragraph_format.space_after=Pt(6)
for name,size,color in [('Title',34,TEXT),('Heading 1',24,TEXT),('Heading 2',15,GREEN),('Heading 3',11,GREEN)]:
    st=styles[name];st.font.name='Liberation Sans';st.font.size=Pt(size);st.font.color.rgb=rgb(color);st.font.bold=True;st.paragraph_format.space_before=Pt(10);st.paragraph_format.space_after=Pt(6)
if 'Eyebrow' not in styles:
    st=styles.add_style('Eyebrow',WD_STYLE_TYPE.PARAGRAPH);st.font.name='Liberation Sans';st.font.size=Pt(8);st.font.bold=True;st.font.color.rgb=rgb(GREEN);st.paragraph_format.space_after=Pt(5)
header=sec.header.paragraphs[0];r=header.add_run('POWERCHAIN  |  CRISIS CAPITAL NETWORK');set_font(r,8,GREEN,True)
page_num(sec.footer.paragraphs[0])

# 1 cover
p=doc.add_paragraph();p.paragraph_format.space_after=Pt(6);p.add_run().add_picture(str(LOGO),width=Inches(1.05))
p=doc.add_paragraph();p.paragraph_format.space_after=Pt(0);r=p.add_run('Power');set_font(r,22,GREEN,True);r=p.add_run('Chain');set_font(r,22,GREEN,False)
eyebrow('DISASTER RELIEF  |  CRISIS CAPITAL NETWORK')
for _ in range(2): doc.add_paragraph('')
p=doc.add_paragraph();r=p.add_run('Verified capital infrastructure\nfor disaster response, recovery\nand measurable impact.');set_font(r,31,TEXT,True);p.paragraph_format.space_after=Pt(18)
p=doc.add_paragraph();r=p.add_run('PWRC Utility Architecture  ·  Solana Settlement  ·  Evidence & Impact Fabric');set_font(r,11,GREEN,True)
p=doc.add_paragraph();p.paragraph_format.space_before=Pt(30);r=p.add_run('Technical Whitepaper\nVersion 1.0.0\n28 August 2026');set_font(r,10,MUTED)
callout('Core principle','PWRC powers utility around the response. Stable assets fund the response. Identity and policy control authority. Cryptographic signers execute value movement. Evidence proves real-world outcomes.')

# 2 exec
newpage();eyebrow('EXECUTIVE SUMMARY');h1('PowerChain connects capital to outcomes')
para('PowerChain Crisis is a capital-coordination and evidence system for verified emergency response. It separates crisis intelligence, stable-asset capital, response execution, trust, settlement and proof of impact so no single wallet, token, AI model or database becomes the entire control plane.',True)
figure('pwrc-utility.png','Figure 1. PWRC is a network utility layer, not crisis donor capital.',6.55)
table(['Challenge','PowerChain control'],[
('Slow / fragmented funding','Solana settlement, checkout, Solana Pay and portable Actions/Blinks.'),('Misallocation risk','Purpose-bound pools, allocation policy, escrow and approval thresholds.'),('Weak transparency','Explicit financial states and public capital-to-impact views.'),('Procurement risk','Supplier controls, explainable integrity signals and evidence.'),('Difficult verification','Evidence Fabric, field verification and Impact Passports.'),('AI overreach','AI may recommend; policy and authorized signers control irreversible actions.')],[2.2,4.55])

# 3 architecture
newpage();eyebrow('ARCHITECTURE');h1('Five connected networks inside CCN')
para('The Crisis Capital Network is a network of bounded domains. The product UI can unify them, but the data and authority models remain separate.')
table(['Network','Primary responsibility'],[
('Crisis','Canonical crisis identity, severity, geography, lifecycle and verified need.'),('Capital','Pools, capital gap, allocations, escrow and funding intents.'),('Response','Organizations, suppliers, procurement, logistics and field operations.'),('Trust','Identity, evidence, policy, risk, verification, attestations and audit.'),('Settlement','Solana programs/tokens, Actions/Blinks and provider-specific cross-chain settlement.')],[1.45,5.3])
h2('Truth layers');bullets(['Solana is authoritative for blockchain settlement and program/token state.','PostgreSQL/Supabase is authoritative for operational workflows, organizations, policies, evidence metadata and audit indexes.','Private object storage holds raw sensitive evidence.','The event engine normalizes authoritative chain/application events for ledger, risk, notifications and analytics.'])
callout('Non-goal','PowerChain does not attempt to put every photo, medical record, personal identity, AI context or large document on-chain.','blue')

# 4 capital lifecycle
newpage();eyebrow('CAPITAL CONTROL');h1('Financial states remain distinct')
figure('source-of-truth.png','Figure 2. Settlement is advanced only after authoritative reconciliation.',6.6)
table(['State','Meaning'],[
('NEED','Verified capital requirement.'),('COMMITTED','Capital promised/reserved by an accepted source.'),('AVAILABLE','Settled capital available to allocate.'),('ALLOCATED','Budget assigned to a permitted purpose/recipient class.'),('ESCROWED','Purpose-bound capital reserved under release conditions.'),('RELEASED','Authorized blockchain/payment movement completed.'),('DELIVERED','Operational delivery record exists.'),('VERIFIED','Required evidence/verification threshold is satisfied.'),('RECONCILED','Financial and operational records agree.')],[1.4,5.35])

# 5 fees
newpage();eyebrow('FEE MODEL');h1('Show the entire fee stack before authorization')
figure('fee-stack.png','Figure 3. Material fee and policy fields are bound into the canonical quote.',6.55)
para('The contributor checkout shows principal and each applicable service, token-transfer, network or provider fee separately. The 5% successful-funding commission is not a contributor checkout charge: where a capital-pool policy enables it, it is charged from successful pool proceeds only after the authoritative success threshold is satisfied.')
h2('Successful-funding commission');table(['Destination','Share of successful pool proceeds'],[('Community Treasury','2 percentage points'),('Ecosystem & Development','3 percentage points'),('Total','5 percentage points')],[3.4,3.35])
callout('Atomic settlement','Where the asset/network supports it, the 2% and 3% legs should be signed in one transaction and reconciled before the commission record becomes CONFIRMED.')

# 6 pwrc
newpage();eyebrow('PWRC 1.0');h1('PowerChain Network Utility')
para('PWRC is the shared utility asset for PowerChain access, AI/API consumption, participation, reviewed incentives and economic governance. A PWRC balance does not confer crisis-treasury or emergency authority.',True)
table(['Utility category','Examples'],[
('ACCESS','Entitlements and service tiers.'),('COMPUTE','AI workloads and bounded agent utility budgets.'),('SERVICES','Reports, data products and selected platform services.'),('PARTICIPATION','Programs and ecosystem coordination.'),('GOVERNANCE','Economic governance; not emergency treasury authority.'),('INCENTIVES','Evidence-backed community/ecosystem rewards.'),('REPUTATION','One input to eligibility/trust, never the sole score.'),('NETWORK','API and machine-service metering.')],[1.5,5.25])
h2('Power Units');para('Power Units are an internal metering abstraction. The reference planning conversion is configurable; service pricing can change without altering token decimals or crisis-capital accounting.')

# 7 token factory
newpage();eyebrow('TOKEN FACTORY');h1('One factory with safer issuer profiles')
figure('token-factory.png','Figure 4. Project, company and government assets share the same factory but use different policy defaults.',6.55)
table(['Profile','Safer default posture'],[
('PROJECT','Capped mint, verified issuer, purpose/metadata disclosure.'),('COMPANY','Verified organization, multisig/HSM policy, explicit treasury/mint separation.'),('GOVERNMENT','Restricted-transfer review by default and stronger change control.'),('RESERVE BACKED','Reserve asset, ratio, attestation, redemption and controlled mint/burn policy required.')],[1.55,5.2])
callout('Mint secret boundary','The backend should prepare and verify Token-2022 transactions, but it should not receive or persist the mint private key. A one-billion-unit value is a planning cap, not an automatic mint.')

# 8 cross-chain
newpage();eyebrow('CROSS-CHAIN');h1('Provider-specific routes fail closed')
figure('cross-chain.png','Figure 5. CCTP, CCIP and wPWRC use different trust and reconciliation models.',6.55)
table(['Route','Purpose','Activation rule'],[
('CCTP V2','Native USDC cross-chain movement.','Domains/version/fees/status verified at execution time.'),('CCIP','Approved messaging/token lanes.','Router/pool/lane, allowlists, rate limits and audit required.'),('wPWRC / Sui','Planned 1:1 PWRC utility representation.','TBA until two-way bridge audit, custody and supply reconciliation are complete.')],[1.35,2.5,2.9])
h2('Conservation invariant');para('Outstanding wPWRC must never exceed verified backing PWRC. Any backing deficit or unknown source-chain state should pause new wrapped minting until reconciliation resolves the discrepancy.')

# 9 security
newpage();eyebrow('SECURITY CONTROL PLANE');h1('Identity, policy and signer authority remain separate')
para('Authentication is not authorization. Token ownership is not authorization. AI output is not authorization.')
table(['Layer','Control'],[
('Identity','Authenticated user/wallet/service identity with replay-resistant challenge/session handling.'),('Organization','Tenant and verified organization membership.'),('Role / permission','Least-privilege operating scope and separation of duties.'),('Policy version','Machine-readable limits, allowed purpose/recipient/asset/network and approvals.'),('Risk','Explainable rules/anomaly signals; high-risk paths may hold/review.'),('Approval','Human or governance threshold appropriate to the action.'),('Signer','Wallet, multisig, HSM/MPC or approved signing infrastructure.'),('Reconciliation','Authoritative blockchain/provider verification before state advances.')],[1.45,5.3])
callout('Hard invariant','No AI agent, client response, PWRC balance or static configuration can bypass treasury policy or cryptographic signing.','red')

# 10 rewards
newpage();eyebrow('COMMUNITY REWARDS');h1('Bounded rewards for verified contribution')
figure('community-rewards.png','Figure 6. Rewards are evidence- and policy-reviewed, not passive-holding emissions.',6.55)
h2('Epoch budget model');bullets(['budget per epoch;','per-subject and per-organization caps;','evidence IDs and duplicate/replay checks;','conflict-of-interest and anomaly review;','challenge/reversal procedure;','separate reward and crisis-capital ledgers.'])
callout('Authority boundary','Reward eligibility does not grant treasury authority, procurement access, verifier status or an automatic reputation upgrade.')

# 11 evidence
newpage();eyebrow('EVIDENCE & IMPACT');h1('A transaction is not an outcome')
para('PowerChain closes the lifecycle with field evidence, verification and outcome records. Sensitive raw evidence remains private/off-chain by default; public transparency uses approved metadata, hashes, attestations and human-readable records.')
table(['Stage','Control'],[
('Upload intent','Organization/crisis/subject binding, MIME/size policy and visibility class.'),('Storage','Private evidence bucket by default; public media uses a separate class.'),('Provenance','Server verifies stored object hash before evidence becomes AVAILABLE.'),('Verification','Verifier/policy records VERIFIED, PARTIAL, PENDING, CONFLICTING, EXPIRED or REJECTED.'),('Attestation','Optional on-chain hash/reference without publishing private content.'),('Impact','Outcome records link back to allocation, procurement, delivery and evidence.'),('Passport','Public or contributor-facing trace from source to verified impact.')],[1.45,5.3])

# 12 UX
newpage();eyebrow('UX / UI');h1('Public simplicity, operator control, auditor evidence')
table(['Audience','Primary experience'],[
('Public','What happened? Can I trust this? Where is capital going? What has been verified?'),('Contributor','Amount, purpose, asset, network, fee stack, destination, authorization and receipt.'),('Operator','Alerts, approvals, blocked capital, procurement exceptions and verification backlog.'),('Treasury','Balances by state, allocations, escrow, releases, signers and reconciliation.'),('Auditor','Source, policy, approval, transaction, recipient, document, evidence and impact chain.')],[1.45,5.3])
h2('Visual language');bullets(['White / light gray / near-black / dark forest green.','No neon or decorative crypto effects.','Color plus text/icon for every status.','Source and freshness labels on live data.','Technical transaction details are available but not forced into the primary public flow.','Mobile prioritizes understand -> trust -> fund -> track -> impact.'])

# 13 demo Nepal
newpage();eyebrow('CANONICAL DEMO FIXTURE');h1('Nepal Flood Response - demo only')
callout('Demo data','The figures on this page are product fixtures. They are not real fundraising totals, verified humanitarian claims or live field outcomes.','blue')
table(['Metric','Demo value'],[('Raised','$4.28M'),('Allocated','$3.71M'),('Verified Impact','$2.94M'),('In Escrow','$770K'),('Verification rate','72%')],[3.2,3.5])
h2('Purpose breakdown');table(['Purpose','Demo allocation'],[('Water','$1.20M'),('Medical','$1.70M'),('Shelter','$900K'),('Logistics','$480K')],[3.2,3.5])
para('Mainnet/public production pages must either display authoritative live data with source identity/freshness or show an explicit unavailable/degraded state. Demo fixtures must never silently substitute for live crisis data.')

# 14 data/oracles
newpage();eyebrow('DATA / ORACLES');h1('Freshness and provenance are first-class')
para('Pricing and external data observations carry provider identity, source, publish timestamp, fetch timestamp, freshness and status. The system must not invent a price or treat an external report as a verified crisis fact.')
table(['Provider / layer','Role'],[
('Pyth','Primary market-price observation where configured; freshness/confidence enforced.'),('Birdeye','Corroboration or degraded fallback according to policy.'),('Chainlink','Provider-specific data/CCIP boundary where configured.'),('Helius','Solana RPC/API/webhook event infrastructure; webhook events are normalized/idempotent.'),('PowerChain backend','Policy evaluation, reconciliation, evidence metadata, risk and operational state.')],[1.6,5.15])
h2('Provider state');para('Use LIVE, DEGRADED, DISABLED, TBA or UNAVAILABLE. Static configuration alone does not justify a LIVE label.')

# 15 production
newpage();eyebrow('PRODUCTION READINESS');h1('Fail closed before mainnet')
table(['Gate','Required evidence'],[
('Dependencies','Pinned runtime/package manager, committed lockfile and clean frozen install.'),('Database','Prisma validation/migrations, RLS/storage review, backup/restore test.'),('Programs','Anchor build/tests, deployment IDs, authority/custody policy and independent review.'),('Financial','Treasury policy, transaction simulation, exact reconciliation and fee disclosure.'),('Cross-chain','Provider configuration, limits, replay protection, monitoring, incident runbook and audit status.'),('Token factory','Issuer approval, metadata, signer flow, reserve review and extension compatibility.'),('AI','Model/tool permission evaluation, prompt-injection testing and audit trail.'),('Operations','Monitoring, DR, event replay, degraded/offline modes and on-call procedures.')],[1.45,5.3])
callout('Unknown execution','If the final blockchain/provider result cannot be established, preserve EXECUTION_UNKNOWN and reconcile later. Never convert uncertainty into success.')

# 16 roadmap
newpage();eyebrow('STRATEGIC ROADMAP');h1('Activation sequence - version remains 1.0.0')
for title,body in [
('01  Core financial close','Audit/deploy core Solana programs, real treasury registry, USDC/Solana Pay and authoritative reconciliation.'),
('02  Evidence close','Production storage/RLS, verifier workflow, impact passport and public transparency.'),
('03  PWRC utility close','Power Units, enterprise utility accounts, AI/API budgets and usage receipts.'),
('04  Token factory close','Issuer policy, secure mint signing, reserve-backed controls and Token-2022 compatibility matrix.'),
('05  Cross-chain close','CCTP production routes, audited CCIP lanes and only then wPWRC/Sui activation.'),
('06  Community reward pilot','Epoch budgets, evidence policy, anti-gaming, review and reversals.'),
('07  Partnership network','Scoped partner credentials, verified organizations and integration contracts.'),
('08  Operational SLOs','RPC/webhook/queue/reconciliation/evidence backlog targets and incident dashboards.')]:
    h2(title);para(body)

# 17 conclusion / references
newpage();eyebrow('CONCLUSION & REFERENCES');h1('Speed without uncontrolled authority')
para('The architectural goal is not to put donations on a blockchain. It is to make emergency capital traceable, policy-controlled and verifiable from crisis identity to real-world outcome.',True)
callout('PowerChain promise','Move capital where it matters. Control how it is used. Prove what it achieved.')
h2('Core invariant');para('PWRC powers the network around the response. Stable assets fund the response. Identity and policy control authority. Cryptographic signers authorize value movement. Evidence proves real-world outcomes.')
h2('Primary technical references')
for label,url in [
('Solana Tokenization','https://solana.com/docs/tokenization'),('Solana Confidential Transfer integration','https://solana.com/docs/tokens/extensions/confidential-transfer/integration-guide'),('Circle CCTP','https://developers.circle.com/cctp'),('Chainlink CCIP','https://docs.chain.link/ccip'),('Pyth Network','https://docs.pyth.network/'),('Helius','https://www.helius.dev/docs')]:
    p=doc.add_paragraph();r=p.add_run(label+': ');set_font(r,9,GREEN,True);r=p.add_run(url);set_font(r,8.5,BLUE)
para('Technical status, supported routes, package IDs, mint addresses, fees, provider capabilities and audit/deployment state must be verified against the running environment. This whitepaper is not financial, legal, investment or tax advice.')

props=doc.core_properties;props.title='PowerChain Crisis Capital Network Whitepaper v1.0.0';props.subject='PWRC Utility Architecture, Solana settlement, token factory, cross-chain security, community rewards and proof of impact';props.author='PowerChain Protocol';props.keywords='PowerChain, Crisis Capital Network, PWRC, Solana, disaster relief, CCTP, CCIP, Token-2022, impact';props.comments='Architecture whitepaper for PowerChain Crisis v1.0.0.'
doc.save(OUT);print(OUT)
